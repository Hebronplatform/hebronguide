#!/usr/bin/env python3
"""
HebronGuide CLAUDE.md → Word 문서 변환기
기록 날짜와 시간을 정확히 기입하여 역사·변천 자료로 보존
"""

import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pytz

# ── 시간 설정 ────────────────────────────────────────────────────────────────
SEOUL_TZ    = pytz.timezone("Asia/Seoul")
SEATTLE_TZ  = pytz.timezone("America/Los_Angeles")
now_utc     = datetime.now(pytz.utc)
now_seattle = now_utc.astimezone(SEATTLE_TZ)
now_seoul   = now_utc.astimezone(SEOUL_TZ)

RECORD_TIME_SEATTLE = now_seattle.strftime("%Y년 %m월 %d일 %H:%M:%S (시애틀)")
RECORD_TIME_SEOUL   = now_seoul.strftime("%Y년 %m월 %d일 %H:%M:%S (한국)")
FILE_TIMESTAMP      = now_seattle.strftime("%Y%m%d_%H%M%S")

CLAUDE_MD_PATH = "CLAUDE.md"
OUTPUT_DIR     = "docs"
OUTPUT_PATH    = f"{OUTPUT_DIR}/HebronGuide_CLAUDE_{FILE_TIMESTAMP}.docx"

import os
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── 헬퍼 함수 ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """테이블 셀 배경색 설정"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_horizontal_line(doc):
    """수평선 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_cover_page(doc):
    """표지 페이지"""
    # 상단 여백
    for _ in range(4):
        doc.add_paragraph()

    # 메인 타이틀
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HebronGuide")
    run.font.size    = Pt(36)
    run.font.bold    = True
    run.font.color.rgb = RGBColor(0xC9, 0xA2, 0x27)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("프로젝트 운영 원칙 & 변천 기록")
    run2.font.size  = Pt(20)
    run2.font.bold  = True
    run2.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

    doc.add_paragraph()

    # 부제
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = desc.add_run("재외 한인 700만을 위한 세계 최고 도시 가이드 플랫폼\nCLAUDE.md 공식 문서화 기록")
    run3.font.size  = Pt(13)
    run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for _ in range(3):
        doc.add_paragraph()

    add_horizontal_line(doc)

    # 기록 시간
    doc.add_paragraph()
    time_p = doc.add_paragraph()
    time_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = time_p.add_run(f"기록 일시\n")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

    r2 = time_p.add_run(f"🇺🇸 {RECORD_TIME_SEATTLE}\n🇰🇷 {RECORD_TIME_SEOUL}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = src.add_run("출처: CLAUDE.md (HebronGuide 프로젝트 공식 운영 문서)\nhttps://github.com/kchurch911/hebronguide")
    r3.font.size  = Pt(9)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def parse_and_write_content(doc, md_text):
    """CLAUDE.md 파싱 후 Word 문서로 변환"""

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines    = []
    in_table      = False
    table_rows    = []

    while i < len(lines):
        line = lines[i]

        # ── 코드 블록 ──
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines    = []
                i += 1
                continue
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name  = "Courier New"
                    run.font.size  = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
                i += 1
                continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── 테이블 ──
        if line.strip().startswith("|"):
            if not in_table:
                in_table   = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                # 구분선 행(---) 제거
                data_rows = [r for r in table_rows if not all(set(c) <= {"-", " ", ":"} for c in r)]
                if len(data_rows) >= 2:
                    cols = len(data_rows[0])
                    tbl  = doc.add_table(rows=len(data_rows), cols=cols)
                    tbl.style = "Table Grid"
                    for ri, row_data in enumerate(data_rows):
                        for ci, cell_text in enumerate(row_data[:cols]):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                set_cell_bg(cell, "1a3a6b")
                    doc.add_paragraph()

        # ── 제목 ──
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.runs[0].font.color.rgb = RGBColor(0xC9, 0xA2, 0x27)
            p.paragraph_format.space_before = Pt(16)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
            p.paragraph_format.space_before = Pt(12)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.paragraph_format.space_before = Pt(8)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)

        # ── 수평선 ──
        elif line.strip() == "---":
            add_horizontal_line(doc)

        # ── 인용 블록 ──
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(line[2:].strip())
            run.font.italic    = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x99)

        # ── 목록 ──
        elif line.startswith("- ") or line.startswith("  - "):
            indent = 1 if line.startswith("  - ") else 0
            text = line.lstrip("- ").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(indent * 0.5)
            _add_inline_formatting(p, text)

        # ── 번호 목록 ──
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text)

        # ── 빈 줄 ──
        elif line.strip() == "":
            pass  # 빈 줄은 그냥 넘김

        # ── 일반 텍스트 ──
        else:
            if line.strip():
                p = doc.add_paragraph()
                _add_inline_formatting(p, line.strip())

        i += 1


def _add_inline_formatting(p, text):
    """굵게(**), 인라인코드(`) 등 인라인 포맷 처리"""
    # **bold** 와 `code` 파싱
    pattern = re.compile(r"(\*\*.*?\*\*|`[^`]+`)")
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                p.add_run(part)


def add_footer_timestamp(doc):
    """각 섹션에 기록 시각 푸터 추가"""
    doc.add_page_break()
    add_horizontal_line(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run(
        f"본 문서는 HebronGuide CLAUDE.md를 기반으로 자동 생성되었습니다.\n"
        f"생성 일시 — 🇺🇸 {RECORD_TIME_SEATTLE}  |  🇰🇷 {RECORD_TIME_SEOUL}\n"
        f"© HebronGuide / 지구촌교회 글로벌 한인 사역"
    )
    r.font.size  = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── 메인 ─────────────────────────────────────────────────────────────────────

def main():
    print(f"📄 Word 문서 생성 시작...")
    print(f"   시애틀: {RECORD_TIME_SEATTLE}")
    print(f"   한국:   {RECORD_TIME_SEOUL}")

    with open(CLAUDE_MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # 기본 폰트
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # 표지
    add_cover_page(doc)

    # 본문
    parse_and_write_content(doc, md_text)

    # 푸터
    add_footer_timestamp(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ 완료: {OUTPUT_PATH}")
    print(f"   파일 크기: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
